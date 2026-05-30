"""Generate Pandoc reference DOCX templates used to style ATS-friendly CVs.

Pandoc only copies STYLES from the reference doc — it doesn't copy
content. So this script programmatically creates three sibling DOCX
templates with distinct ATS-safe style profiles:

  - plain   — Arial 11pt, 0.55" margins, 18pt black H1, ALL-CAPS H2 with
              full-width hairline border, very neutral. Safest for older
              parsers / pure keyword scrapers.
  - classic — Times New Roman 11pt, 0.7" margins, centered 22pt H1 with
              letter-spacing, Title-Case H2 (no all-caps). The "executive
              search firm" look — leans more human-reader than ATS-only.
  - modern  — Calibri 11pt, 0.6"x0.7" margins, navy 24pt H1 + 13pt
              ALL-CAPS H2 with bottom border. The default JobHunter look
              — same styles the legacy `cv-ats-template.docx` shipped.

Why three? Different applications benefit from different aesthetics —
boring/plain wins on broken Workday parsers, modern wins on FAANG/agency
roles where a recruiter actually opens the file. The downstream renderer
picks one per application.

ATS-parser-friendly choices, common to all three:
  - Single-column layout — multi-column resumes routinely fail parsing
  - Calibri / Arial / Times New Roman — fonts ATS engines whitelist
  - No graphic shapes, text boxes — parsers strip them and lose content
  - Date ranges as plain text inline, not in a side column
  - List Bullet style only (no nested numbered lists, no manual bullets)

Usage (from repo root):

    python backend/scripts/generate_cv_template.py

Re-run after editing this file to refresh templates. Outputs go to
`backend/templates/cv-template-{plain,classic,modern}.docx` plus a
back-compat copy at `backend/templates/cv-ats-template.docx` (modern
variant) — the latter is consumed by code paths that haven't been
migrated to per-template lookup yet.
"""
from __future__ import annotations

import shutil
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

TEMPLATES_DIR = Path(__file__).resolve().parents[1] / "templates"
LEGACY_OUT_PATH = TEMPLATES_DIR / "cv-ats-template.docx"


@dataclass(frozen=True)
class TemplateSpec:
    """Style profile for a single reference DOCX.

    Field-by-field rationale:
      - font_name: ATS whitelisted (Arial / Times New Roman / Calibri).
      - margins_in: tighter than Word default (1.0") to fit ~10% more
        content; 0.55-0.7" is the sweet spot.
      - h1_*: candidate name styling (size, alignment, color).
      - h2_*: section header styling — ALL CAPS toggle and color.
      - h3_size_pt: role/entry header size. Always bold.
      - body_color: Normal + List Bullet text color.
      - h2_border_size: bottom-border weight in 1/8 pt units (Word's
        `w:sz` value: 6 = 0.75pt, 8 = 1pt).
      - bullet_space_after_pt: tightness of bullet stack — classic uses
        1pt for an executive-tight feel, plain/modern use 2pt.
      - h1_letter_spacing_pt: positive = expanded tracking (used on the
        classic centered name to give it a "monogram" feel).
    """

    font_name: str
    margins_in: float
    margins_in_lr: float
    h1_size_pt: float
    h1_centered: bool
    h1_color: tuple[int, int, int]
    h1_letter_spacing_pt: float
    h2_size_pt: float
    h2_all_caps: bool
    h2_color: tuple[int, int, int]
    h2_border_size: int
    h3_size_pt: float
    body_color: tuple[int, int, int]
    bullet_space_after_pt: float


# Specs derived from the Phase A brief table. Do not collapse fields
# across templates "to deduplicate" — the divergence is the point.
TEMPLATE_SPECS: dict[str, TemplateSpec] = {
    "plain": TemplateSpec(
        font_name="Arial",
        margins_in=0.55,
        margins_in_lr=0.55,
        h1_size_pt=18,
        h1_centered=False,
        h1_color=(0, 0, 0),
        h1_letter_spacing_pt=0,
        h2_size_pt=11.5,
        h2_all_caps=True,
        h2_color=(0, 0, 0),
        h2_border_size=8,  # 1pt full-width
        h3_size_pt=11,
        body_color=(0, 0, 0),
        bullet_space_after_pt=2,
    ),
    "classic": TemplateSpec(
        font_name="Times New Roman",
        margins_in=0.7,
        margins_in_lr=0.7,
        h1_size_pt=22,
        h1_centered=True,
        h1_color=(26, 26, 26),
        h1_letter_spacing_pt=1,
        h2_size_pt=12,
        h2_all_caps=False,
        h2_color=(26, 26, 26),
        h2_border_size=6,  # 0.75pt
        h3_size_pt=11.5,
        body_color=(26, 26, 26),
        bullet_space_after_pt=1,
    ),
    "modern": TemplateSpec(
        font_name="Calibri",
        margins_in=0.6,
        margins_in_lr=0.7,
        h1_size_pt=24,
        h1_centered=False,
        h1_color=(20, 30, 50),
        h1_letter_spacing_pt=0,
        h2_size_pt=13,
        h2_all_caps=True,
        h2_color=(20, 30, 50),
        h2_border_size=6,  # 0.75pt
        h3_size_pt=12,
        body_color=(0, 0, 0),
        bullet_space_after_pt=2,
    ),
}


def _set_paragraph_border(paragraph, *, bottom: bool = False, size: int = 6):
    """Add a hairline border below the paragraph.

    `size` is Word's `w:sz` (eighths of a point): 6 = 0.75pt, 8 = 1pt."""
    if not bottom:
        return
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom_bdr = OxmlElement("w:bottom")
    bottom_bdr.set(qn("w:val"), "single")
    bottom_bdr.set(qn("w:sz"), str(size))
    bottom_bdr.set(qn("w:space"), "1")
    bottom_bdr.set(qn("w:color"), "808080")
    p_bdr.append(bottom_bdr)
    p_pr.append(p_bdr)


def _configure_style(
    doc,
    name: str,
    *,
    font_name: str = "Calibri",
    font_size_pt: float | None = None,
    bold: bool = False,
    italic: bool = False,
    color: tuple[int, int, int] | None = None,
    space_before_pt: float = 0,
    space_after_pt: float = 0,
    all_caps: bool = False,
    keep_with_next: bool = False,
    centered: bool = False,
    letter_spacing_pt: float = 0,
):
    """Update a built-in style by name. Pandoc maps:
        Heading 1 -> # in markdown
        Heading 2 -> ##
        Heading 3 -> ###
        Normal    -> body text
        List Bullet -> - / * lists
    Keeping the names matches Pandoc's expectations exactly."""
    style = doc.styles[name]
    font = style.font
    font.name = font_name
    if font_size_pt is not None:
        font.size = Pt(font_size_pt)
    font.bold = bold
    font.italic = italic
    if color is not None:
        font.color.rgb = RGBColor(*color)

    pf = style.paragraph_format
    pf.space_before = Pt(space_before_pt)
    pf.space_after = Pt(space_after_pt)
    pf.keep_with_next = keep_with_next
    if centered:
        pf.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    rpr = style.element.get_or_add_rPr()

    if all_caps:
        caps = OxmlElement("w:caps")
        caps.set(qn("w:val"), "true")
        rpr.append(caps)

    if letter_spacing_pt:
        # Word's `w:spacing` character property uses 1/20 pt units —
        # 1pt of tracking is "20" in OOXML. Used on the classic H1.
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:val"), str(int(letter_spacing_pt * 20)))
        rpr.append(spacing)


def _build_template(name: str, spec: TemplateSpec) -> Path:
    """Build one reference DOCX for the given spec and write it to
    `backend/templates/cv-template-{name}.docx`. Returns the output path."""
    out_path = TEMPLATES_DIR / f"cv-template-{name}.docx"
    out_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(spec.margins_in)
        section.bottom_margin = Inches(spec.margins_in)
        section.left_margin = Inches(spec.margins_in_lr)
        section.right_margin = Inches(spec.margins_in_lr)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)

    _configure_style(
        doc,
        "Normal",
        font_name=spec.font_name,
        font_size_pt=11,
        color=spec.body_color,
        space_after_pt=4,
    )

    _configure_style(
        doc,
        "Heading 1",
        font_name=spec.font_name,
        font_size_pt=spec.h1_size_pt,
        bold=True,
        color=spec.h1_color,
        space_before_pt=0,
        space_after_pt=2,
        keep_with_next=True,
        centered=spec.h1_centered,
        letter_spacing_pt=spec.h1_letter_spacing_pt,
    )

    _configure_style(
        doc,
        "Heading 2",
        font_name=spec.font_name,
        font_size_pt=spec.h2_size_pt,
        bold=True,
        color=spec.h2_color,
        space_before_pt=12,
        space_after_pt=4,
        all_caps=spec.h2_all_caps,
        keep_with_next=True,
    )

    _configure_style(
        doc,
        "Heading 3",
        font_name=spec.font_name,
        font_size_pt=spec.h3_size_pt,
        bold=True,
        color=spec.body_color,
        space_before_pt=8,
        space_after_pt=2,
        keep_with_next=True,
    )

    _configure_style(
        doc,
        "List Bullet",
        font_name=spec.font_name,
        font_size_pt=11,
        color=spec.body_color,
        space_before_pt=0,
        space_after_pt=spec.bullet_space_after_pt,
    )

    if "Quote" in [s.name for s in doc.styles]:
        _configure_style(
            doc,
            "Quote",
            font_name=spec.font_name,
            font_size_pt=10,
            italic=True,
            color=(80, 80, 80),
            space_before_pt=2,
            space_after_pt=4,
        )

    # Sample content — Pandoc copies STYLES only from the reference,
    # so this is just a designer-friendly preview when the file is
    # opened in Word directly.
    h1 = doc.add_paragraph("Your Name", style="Heading 1")
    if spec.h1_centered:
        h1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    contact = doc.add_paragraph(
        "your.email@example.com · +00 000 000 000 · City, Country"
    )
    contact.runs[0].font.size = Pt(10)
    if spec.h1_centered:
        contact.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    summary_label = "Summary" if not spec.h2_all_caps else "SUMMARY"
    h2 = doc.add_paragraph(summary_label, style="Heading 2")
    _set_paragraph_border(h2, bottom=True, size=spec.h2_border_size)
    doc.add_paragraph(
        "Two-to-three sentence professional summary highlighting your most "
        "relevant experience and the value you bring to the role."
    )

    exp_label = "Experience" if not spec.h2_all_caps else "EXPERIENCE"
    h2 = doc.add_paragraph(exp_label, style="Heading 2")
    _set_paragraph_border(h2, bottom=True, size=spec.h2_border_size)
    doc.add_paragraph("Job Title · Company Name", style="Heading 3")
    meta = doc.add_paragraph("Jan 2024 – Present · Location")
    meta.runs[0].italic = True
    meta.runs[0].font.size = Pt(10)
    meta.runs[0].font.color.rgb = RGBColor(80, 80, 80)
    doc.add_paragraph(
        "One or two sentences summarizing the role and your scope."
    )
    doc.add_paragraph("Quantified achievement bullet 1", style="List Bullet")
    doc.add_paragraph("Quantified achievement bullet 2", style="List Bullet")

    doc.save(out_path)
    return out_path


def main() -> int:
    TEMPLATES_DIR.mkdir(parents=True, exist_ok=True)

    modern_path: Path | None = None
    for name, spec in TEMPLATE_SPECS.items():
        out_path = _build_template(name, spec)
        print(f"wrote {out_path} ({out_path.stat().st_size} bytes)")
        if name == "modern":
            modern_path = out_path

    # Back-compat shim — the existing API paths still load the legacy
    # filename. Phase E will drop both this copy and the consumer code.
    if modern_path is not None:
        shutil.copyfile(modern_path, LEGACY_OUT_PATH)
        print(
            f"wrote {LEGACY_OUT_PATH} ({LEGACY_OUT_PATH.stat().st_size} bytes)"
            " [legacy alias — copy of modern]"
        )

    return 0


if __name__ == "__main__":
    raise SystemExit(main())
